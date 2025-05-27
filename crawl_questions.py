from selenium import webdriver
from selenium.webdriver.edge.service import Service
from bs4 import BeautifulSoup
from docx import Document
import time
import datetime
import re
import os

def fetch_questions(url):
    # Set up Edge WebDriver
    service = Service("D:\\OneDrive - HUCE\\2_Code\\Python\\Craw_data\\msedgedriver.exe")
    driver = webdriver.Edge(service=service)
    driver.get(url)
    time.sleep(5)

    # Parse the page source with BeautifulSoup
    soup = BeautifulSoup(driver.page_source, 'html.parser')
    driver.quit()

    # Locate questions and answers
    content = soup.find('div', class_='entry-content')
    if not content:
        print("Unable to find the main content.")
        return []
    
    questions = []
    current_question = None

    for element in content.find_all(['p', 'div']):
        text = element.get_text(strip=True)
        # Remove "View Answer" text
        text = text.replace('View Answer', '')
        
        # Identify questions starting with numbers
        if element.name == 'p' and text.strip().startswith(tuple(str(i) + "." for i in range(1, 100))):
            # Split the text by options (a), b), c), d))
            parts = re.split(r'([a-d]\))', text)
            
            # First part is the question
            question_text = parts[0].strip()
            
            # Process options
            options = []
            for i in range(1, len(parts)-1, 2):
                if parts[i] and parts[i+1]:  # Make sure we have both letter and text
                    option_letter = parts[i][0].upper()  # Convert 'a' to 'A'
                    option_text = parts[i+1].strip()
                    options.append((option_letter, option_text))
            current_question = (question_text, options)
        
        elif element.name == 'div' and 'collapseomatic_content' in element.get('class', []):
            answer_text = element.get_text(strip=True)
            # Extract just the letter from "Answer: X" format
            answer_match = re.search(r'Answer:\s*([a-d])', answer_text, re.IGNORECASE)
            if answer_match:
                answer_text = answer_match.group(1)
            if current_question:
                questions.append((*current_question, answer_text))
                current_question = None

    return questions

def create_docx(questions, filename):
    doc = Document()
    
    for i, (question, options, answer) in enumerate(questions, 1):
        # Add question as Heading 3
        # Remove the question number if it exists at the start
        question_clean = re.sub(r'^\d+\.\s*', '', question)
        question_text = f"Câu {i}: {question_clean}"
        heading = doc.add_heading(level=3)
        heading.add_run(question_text)
        
        # Add options with correct answer in bold, each on a new line
        for option_letter, option_text in options:
            paragraph = doc.add_paragraph()
            # Compare option letter with answer letter (case-insensitive)
            if option_letter.upper() == answer.upper():
                # Make the entire option (letter and text) bold for correct answer
                run = paragraph.add_run(f"{option_letter}. {option_text}")
                run.bold = True
            else:
                # Regular formatting for incorrect answers
                paragraph.add_run(f"{option_letter}. {option_text}")
        
        # Add a blank line after each question
        doc.add_paragraph()

    doc.save(filename)

def main():
    input_file = 'urls.txt'
    if not os.path.exists(input_file):
        print(f"Input file '{input_file}' not found.")
        return

    with open(input_file, 'r', encoding='utf-8') as f:
        urls = [line.strip() for line in f if line.strip()]

    for url in urls:
        
        questions = fetch_questions(url)
        if not questions:
            print("No questions found. Please check the HTML structure or URL.")
        else:
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{timestamp}.docx"
            create_docx(questions, filename)
            print(f"Document created successfully with {len(questions)} questions. Filename: {filename}")

if __name__ == '__main__':
    main()